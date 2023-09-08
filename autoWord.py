import csv
import math
from docx import Document
from PyPDF2 import PdfWriter, PdfReader
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from datetime import date
import sys
import time
import os

ros_path = "C:\\Users\\Pearson\\Downloads\\"
abs_path = "C:\\Users\\Pearson\\Desktop\\Pearson Vue\\"
rel_path = "C:\\Users\\Pearson\\PycharmProjects\\Python Automation\\"
proctor  = "Darren Enos"

# Progress bar for powershell
def progress(count, total, status=''):
    bar_len = 60
    filled_len = int(round(bar_len * count / float(total)))

    percents = round(100.0 * count / float(total), 1)
    bar = '=' * filled_len + '-' * (bar_len - filled_len)

    sys.stdout.write('[%s] %s%s ...%s\r' % (bar, percents, '%', status))
    sys.stdout.flush()
# Fill PDF using [X,Y] coordinates
def fill_pdf(roster,admin,name,page):
    packet = io.BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    nameYstart = 728
    codeYstart = 710

    can.drawString(540,763,str(page[0]))
    can.drawString(572,763,str(page[1]))
    can.drawString(515,749,date.today().strftime("%m/%d/%y"))

    for i in range(len(roster)):
        can.setFont("Helvetica", 14)
        yPos = nameYstart-(i*136)
        #can.drawString(140,yPos,roster[i][0])
        #can.drawString(455,yPos,admin)

        if len(roster[i][1]) > 11:
            can.setFont("Helvetica",8)
        elif len(roster[i][1]) > 9:
            can.setFont("Helvetica", 10)

        #can.drawString(90,codeYstart-(i*136),roster[i][1])

    can.save()

    #move to the beginning of the StringIO buffer
    packet.seek(0)

    # create a new PDF with Reportlab
    new_pdf = PdfReader(packet)
    # read your existing PDF
    existing_pdf = PdfReader(open(rel_path+"roster.pdf", "rb"))
    output = PdfWriter()
    # add the "watermark" (which is the new pdf) on the existing page
    page = existing_pdf.pages[0]
    page.merge_page(new_pdf.pages[0])
    output.add_page(page)
    # finally, write "output" to a real file
    outputStream = open(abs_path+name+".pdf", "wb")
    output.write(outputStream)
    outputStream.close()
# Fill word template with find & replace
def fill_word(doc,name,stime,exam):
    a = 0
    for p in doc.paragraphs:
        a += 1
        for r in p.runs:
            for k in r._r:
                if k.text:
                    if "*Name*" in k.text:
                        k.text = k.text.replace("*Name*", name).strip()
                    #if "{{Name}}" in k.text:
                    #    k.text = k.text.replace("{{Name}}", name).strip()
                    #if "{{Start}}" in k.text:
                    #    k.text = k.text.replace("{{Start}}", stime).strip()
                    if "{{Exam}}" in k.text:
                        k.text = k.text.replace("{{Exam}}", exam).strip()
    return doc
# Returns list of available rosters from [ros_path]
def find_files(filename, search_path):
   result = []

# Wlaking top-down from the root
   for root, dir, files in os.walk(search_path):
      for file in files:
         if filename in file:
            result.append(os.path.join(root, file))

   return result
# Ask user to select correct roster
def check_file(results):
   file = results.pop(0).split("\\")[-1]
   check = input("Found Roster: "+file+"\nUse this file? (Y/n)")
   if check.capitalize() == "Y":
      return file
   elif check.capitalize() == "N":
      if len(results) > 0:
        print("Roster Selected.")
        return check_file(results)
      else:
        print("Could not find a Roster file!")
        quit()


# MAIN

result = check_file(find_files("Pearson VUE","C:/Users/Pearson/Downloads"))

roster_list = []

# Parse CSV
with open(ros_path+result) as csvfile:
    roster = csv.reader(csvfile)
    next(roster) # Skip
    next(roster) # 4
    next(roster) # header
    next(roster) # lines
    for row in roster:
        if len(row) == 0: # blank row indicates end of roster
            break
        roster_list.append(row)

doc_pile = []
doc_count = len(roster_list)
name_list = []
acc_list = []

# Create template docs
for i in range(doc_count):
    doc_pile.append(Document(rel_path+'template.docx'))

# Fill out word documents
for i in range(doc_count):
    progress((i+1),doc_count,"Creating Docs for "+str(i+1)+"/"+str(doc_count)+" testers")

    stime,name,client,exam,length,acc,brk,notes = roster_list[i]

    # Add name if not in list (for counting unique testers)
    if name not in name_list:
        name_list.append(name)
    # Catch any accommodations
    if acc == 1:
        acc_list.append(name)

    if client == 'Microsoft':
        doc = Document(rel_path+'ms-template.docx')
    elif client == 'Pharmacy Technician Certification Board':
        doc = Document(rel_path+'ptcb-template.docx')
    elif client == 'Evaluation Systems':
        doc = Document(rel_path+'es-template.docx')
    elif 'CDA' in client:
        doc = Document(rel_path+'cda-template.docx')
    elif 'ICC' in client:
        doc = Docuemtn(rel_path+'icc-template.docx')
    else:
        doc = doc_pile[i]
    doc = fill_word(doc,name,stime,exam)
    doc.save(abs_path+str(i)+'.docx')



i = 1
pdfs = []
pages = math.ceil(float(len(roster_list)/5))

# Fill out PDFs (1 for every 5 testers)
print("Building PDFs...")
for entry in roster_list:
    pdfs.append([entry[1],entry[3]])
    if len(pdfs) == 5:
        fill_pdf(pdfs,proctor,str(i),[i,pages])
        pdfs = []
        i += 1

# Put remaining testers on last PDF
if len(pdfs) > 0:
    fill_pdf(pdfs,proctor,str(i),[i,pages])

input("You have "+str(len(name_list))+" testers taking "+str(len(roster_list))+" tests. (press ENTER to continue)")

if len(acc_list) > 0:
    print("You have testers with accommodations:")
    for tester in acc_list:
        print(tester)
    input("Press ENTER to continue")
else:
    print("No accommodations found.")